import os
import argparse
import base64
from typing import List, Union, Tuple
from dotenv import load_dotenv

# Color support for cross-platform compatibility
try:
    from colorama import init, Fore, Style
    init(autoreset=True)
    COLORS_AVAILABLE = True
except ImportError:
    # Fallback if colorama is not available
    class Fore:
        GREEN = ""
        YELLOW = ""
        BLUE = ""
        RED = ""
        CYAN = ""
        MAGENTA = ""
        WHITE = ""
        RESET = ""
    
    class Style:
        BRIGHT = ""
        DIM = ""
        NORMAL = ""
        RESET_ALL = ""
    
    COLORS_AVAILABLE = False
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import DirectoryLoader
from langchain_chroma import Chroma
from langchain.schema import HumanMessage, AIMessage, Document

# File processing imports
import pandas as pd
import docx
import chardet
from PyPDF2 import PdfReader

# Web scraping imports
import requests
from bs4 import BeautifulSoup

# Voice and Audio imports
import speech_recognition as sr
import pyttsx3
import threading
import queue
import time
import wave
import pyaudio

# Image processing imports
from PIL import Image, ImageDraw, ImageFont
import matplotlib.pyplot as plt
import numpy as np
import cairosvg

load_dotenv()

def colored_print(text: str, color: str, style: str = ""):
    """Print text with color and style if colors are available"""
    if COLORS_AVAILABLE:
        print(f"{style}{color}{text}{Style.RESET_ALL}")
    else:
        print(text)

# AI Model configurations
AI_MODELS = {
    'gemini': {
        'name': 'Google Gemini',
        'chat_class': 'ChatGoogleGenerativeAI',
        'embedding_class': 'GoogleGenerativeAIEmbeddings',
        'api_key_env': 'GEMINI_API_KEY',
        'default_model': 'gemini-2.0-flash',
        'embedding_model': 'models/embedding-001'
    },
    'chatgpt': {
        'name': 'OpenAI ChatGPT',
        'chat_class': 'ChatOpenAI',
        'embedding_class': 'OpenAIEmbeddings',
        'api_key_env': 'OPENAI_API_KEY',
        'default_model': 'gpt-4o-mini',
        'embedding_model': 'text-embedding-3-small'
    },
    'claude': {
        'name': 'Anthropic Claude',
        'chat_class': 'ChatAnthropic',
        'embedding_class': 'AnthropicEmbeddings',
        'api_key_env': 'ANTHROPIC_API_KEY',
        'default_model': 'claude-3-5-sonnet-20241022',
        'embedding_model': 'claude-3-5-sonnet-20241022'
    },
    'llama': {
        'name': 'Meta Llama',
        'chat_class': 'ChatOllama',
        'embedding_class': 'OllamaEmbeddings',
        'api_key_env': 'OLLAMA_API_KEY',
        'default_model': 'llama3.2:3b',
        'embedding_model': 'llama3.2:3b'
    },
    'gpt4all': {
        'name': 'GPT4All',
        'chat_class': 'ChatOllama',
        'embedding_class': 'OllamaEmbeddings',
        'api_key_env': 'OLLAMA_API_KEY',
        'default_model': 'gpt4all:latest',
        'embedding_model': 'gpt4all:latest'
    },
    'perplexity': {
        'name': 'Perplexity AI',
        'chat_class': 'ChatPerplexity',
        'embedding_class': 'PerplexityEmbeddings',
        'api_key_env': 'PERPLEXITY_API_KEY',
        'default_model': 'llama-3.1-8b-instruct',
        'embedding_model': 'llama-3.1-8b-instruct'
    }
}

def parse_arguments():
    """Parse command line arguments"""
    parser = argparse.ArgumentParser(
        description='AI Chat with Vector DB and Web Search - Multi-Model Support',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python dami_ai.py --model gemini
  python dami_ai.py --model chatgpt --temperature 0.8
  python dami_ai.py --model claude --model-name claude-3-5-haiku-20241022
  python dami_ai.py --model llama --model-name llama3.2:7b
  python dami_ai.py --model gpt4all --model-name gpt4all:latest
  python dami_ai.py --model perplexity --model-name llama-3.1-8b-instruct
  python dami_ai.py --model gemini --verbose-web-search
        """
    )
    
    parser.add_argument(
        '--model', 
        choices=list(AI_MODELS.keys()),
        default='gemini',
        help='AI model to use (default: gemini)'
    )
    
    parser.add_argument(
        '--model-name',
        help='Specific model name/version to use'
    )
    
    parser.add_argument(
        '--temperature',
        type=float,
        default=0.7,
        help='Temperature for AI responses (default: 0.7)'
    )
    
    parser.add_argument(
        '--max-tokens',
        type=int,
        help='Maximum tokens for AI responses'
    )
    

    
    parser.add_argument(
        '--no-web-search',
        action='store_true',
        help='Disable web search functionality'
    )
    
    parser.add_argument(
        '--chunk-size',
        type=int,
        default=1000,
        help='Document chunk size (default: 1000)'
    )
    
    parser.add_argument(
        '--chunk-overlap',
        type=int,
        default=200,
        help='Document chunk overlap (default: 200)'
    )
    
    parser.add_argument(
        '--verbose-web-search',
        action='store_true',
        help='Show detailed web search detection information'
    )
    
    return parser.parse_args()

class AIModelFactory:
    """Factory class to create different AI models"""
    
    def __init__(self, model_type: str, model_name: str = None, temperature: float = 0.7, max_tokens: int = None):
        self.model_type = model_type
        self.model_name = model_name or AI_MODELS[model_type]['default_model']
        self.temperature = temperature
        self.max_tokens = max_tokens
        self.model_config = AI_MODELS[model_type]
        
        # Check API key
        self.api_key = os.getenv(self.model_config['api_key_env'])
        if not self.api_key:
            raise ValueError(f"{self.model_config['api_key_env']} environment variable not set for {self.model_config['name']}")
    
    def create_chat_model(self):
        """Create and return the appropriate chat model"""
        try:
            if self.model_type == 'gemini':
                from langchain_google_genai import ChatGoogleGenerativeAI
                return ChatGoogleGenerativeAI(
                    model=self.model_name,
                    google_api_key=self.api_key,
                    temperature=self.temperature,
                    max_output_tokens=self.max_tokens
                )
            
            elif self.model_type == 'chatgpt':
                from langchain_openai import ChatOpenAI
                return ChatOpenAI(
                    model=self.model_name,
                    openai_api_key=self.api_key,
                    temperature=self.temperature,
                    max_tokens=self.max_tokens
                )
            
            elif self.model_type == 'claude':
                from langchain_anthropic import ChatAnthropic
                return ChatAnthropic(
                    model=self.model_name,
                    anthropic_api_key=self.api_key,
                    temperature=self.temperature,
                    max_tokens=self.max_tokens
                )
            
            elif self.model_type in ['llama', 'gpt4all']:
                from langchain_community.chat_models import ChatOllama
                return ChatOllama(
                    model=self.model_name,
                    temperature=self.temperature,
                    num_predict=self.max_tokens
                )
            
            elif self.model_type == 'perplexity':
                from langchain_community.chat_models import ChatPerplexity
                return ChatPerplexity(
                    model=self.model_name,
                    api_key=self.api_key,
                    temperature=self.temperature,
                    max_tokens=self.max_tokens
                )
            
            else:
                raise ValueError(f"Unsupported model type: {self.model_type}")
                
        except ImportError as e:
            raise ImportError(f"Required package not installed for {self.model_type}: {e}")
    
    def create_embedding_model(self):
        """Create and return the appropriate embedding model"""
        try:
            if self.model_type == 'gemini':
                from langchain_google_genai import GoogleGenerativeAIEmbeddings
                return GoogleGenerativeAIEmbeddings(
                    model=self.model_config['embedding_model'],
                    google_api_key=self.api_key
                )
            
            elif self.model_type == 'chatgpt':
                from langchain_openai import OpenAIEmbeddings
                return OpenAIEmbeddings(
                    model=self.model_config['embedding_model'],
                    openai_api_key=self.api_key
                )
            
            elif self.model_type == 'claude':
                from langchain_anthropic import AnthropicEmbeddings
                return AnthropicEmbeddings(
                    model=self.model_config['embedding_model'],
                    anthropic_api_key=self.api_key
                )
            
            elif self.model_type in ['llama', 'gpt4all']:
                from langchain_community.embeddings import OllamaEmbeddings
                return OllamaEmbeddings(
                    model=self.model_name
                )
            
            elif self.model_type == 'perplexity':
                from langchain_community.embeddings import PerplexityEmbeddings
                return PerplexityEmbeddings(
                    model=self.model_config['embedding_model'],
                    api_key=self.api_key
                )
            
            else:
                raise ValueError(f"Unsupported model type: {self.model_type}")
                
        except ImportError as e:
            raise ImportError(f"Required package not installed for {self.model_type}: {e}")

class WebSearch:
    """Class to handle web search functionality"""
    
    def __init__(self):
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        })
    
    def search_web(self, query: str, max_results: int = 3) -> List[dict]:
        """
        Perform web search using DuckDuckGo (no API key required)
        Returns list of search results with title, snippet, and URL
        """
        try:
            # Using DuckDuckGo search (no API key required)
            search_url = "https://duckduckgo.com/html/"
            params = {
                'q': query,
                'kl': 'us-en'
            }
            
            response = self.session.get(search_url, params=params, timeout=10)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.text, 'html.parser')
            results = []
            
            # Extract search results
            for result in soup.find_all('div', class_='result')[:max_results]:
                title_elem = result.find('a', class_='result__a')
                snippet_elem = result.find('a', class_='result__snippet')
                url_elem = result.find('a', class_='result__url')
                
                if title_elem and snippet_elem:
                    title = title_elem.get_text(strip=True)
                    snippet = snippet_elem.get_text(strip=True)
                    url = url_elem.get_text(strip=True) if url_elem else "No URL available"
                    
                    results.append({
                        'title': title,
                        'snippet': snippet,
                        'url': url
                    })
            
            return results
            
        except Exception as e:
            print(f"Web search error: {e}")
            return []
    
    def get_web_content(self, url: str, max_length: int = 2000) -> str:
        """
        Fetch content from a specific URL
        Returns cleaned text content
        """
        try:
            response = self.session.get(url, timeout=10)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text content
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            # Limit length
            if len(text) > max_length:
                text = text[:max_length] + "..."
            
            return text
            
        except Exception as e:
            return f"Error fetching content: {e}"

class VoiceManager:
    """Manages voice input/output capabilities"""
    
    def __init__(self, enable_voice=True):
        self.enable_voice = enable_voice
        self.recognizer = None
        self.tts_engine = None
        self.audio_queue = queue.Queue()
        self.is_listening = False
        self.wake_word = "hey ai"
        
        if enable_voice:
            self._initialize_voice()
    
    def _initialize_voice(self):
        """Initialize speech recognition and TTS"""
        try:
            # Initialize speech recognition
            self.recognizer = sr.Recognizer()
            self.recognizer.energy_threshold = 4000
            self.recognizer.dynamic_energy_threshold = True
            
            # Initialize text-to-speech
            self.tts_engine = pyttsx3.init()
            self.tts_engine.setProperty('rate', 150)  # Speed of speech
            self.tts_engine.setProperty('volume', 0.8)  # Volume level
            
            # Get available voices and set a good one
            voices = self.tts_engine.getProperty('voices')
            if voices:
                self.tts_engine.setProperty('voice', voices[0].id)
            
            print("✅ Voice features initialized successfully!")
            
        except ImportError as e:
            print(f"⚠️  Voice features disabled: {e}")
            print("💡 Install with: pip install SpeechRecognition pyttsx3 pyaudio")
            self.enable_voice = False
        except Exception as e:
            print(f"⚠️  Voice features disabled: {e}")
            print("💡 Install with: pip install SpeechRecognition pyttsx3 pyaudio")
            self.enable_voice = False
    
    def speak(self, text: str):
        """Convert text to speech"""
        if not self.enable_voice or not self.tts_engine:
            return
        
        try:
            # Run TTS in a separate thread to avoid blocking
            def speak_thread():
                self.tts_engine.say(text)
                self.tts_engine.runAndWait()
            
            thread = threading.Thread(target=speak_thread)
            thread.start()
            
        except Exception as e:
            print(f"❌ TTS error: {e}")
    
    def listen_for_wake_word(self, timeout=5):
        """Listen for wake word and return True if detected"""
        if not self.enable_voice or not self.recognizer:
            return False
        
        try:
            with sr.Microphone() as source:
                print("🎤 Listening for wake word... (say 'hey ai')")
                self.recognizer.adjust_for_ambient_noise(source, duration=0.5)
                
                try:
                    audio = self.recognizer.listen(source, timeout=timeout, phrase_time_limit=3)
                    text = self.recognizer.recognize_google(audio).lower()
                    
                    if self.wake_word in text:
                        print("🎯 Wake word detected!")
                        return True
                    
                except sr.WaitTimeoutError:
                    pass
                except sr.UnknownValueError:
                    pass
                except sr.RequestError as e:
                    print(f"❌ Speech recognition error: {e}")
                    
        except Exception as e:
            print(f"❌ Microphone error: {e}")
        
        return False
    
    def listen_for_input(self, timeout=10):
        """Listen for voice input and return transcribed text"""
        if not self.enable_voice or not self.recognizer:
            return None
        
        try:
            with sr.Microphone() as source:
                print("🎤 Listening for your question...")
                self.recognizer.adjust_for_ambient_noise(source, duration=0.5)
                
                try:
                    audio = self.recognizer.listen(source, timeout=timeout, phrase_time_limit=15)
                    text = self.recognizer.recognize_google(audio)
                    print(f"🎯 Heard: {text}")
                    return text
                    
                except sr.WaitTimeoutError:
                    print("⏰ No speech detected within timeout")
                except sr.UnknownValueError:
                    print("❓ Could not understand speech")
                except sr.RequestError as e:
                    print(f"❌ Speech recognition error: {e}")
                    
        except Exception as e:
            print(f"❌ Microphone error: {e}")
        
        return None
    
    def start_voice_mode(self):
        """Start continuous voice interaction mode"""
        if not self.enable_voice:
            print("❌ Voice features not available")
            return
        
        print("🎤 Voice mode activated! Say 'hey ai' to wake me up.")
        print("💡 Say 'exit voice' to return to text mode.")
        
        while True:
            if self.listen_for_wake_word(timeout=10):
                self.speak("Hello! How can I help you?")
                
                # Listen for question
                question = self.listen_for_input()
                if question:
                    if "exit voice" in question.lower():
                        self.speak("Returning to text mode")
                        break
                    
                    # Return the question for processing
                    return question
                    
                self.speak("I didn't catch that. Please try again.")
        
        return None

class ImageLoader:
    """Class to handle image loading and CLIP embedding generation"""
    
    def __init__(self):
        self.clip_model = None
        self.clip_processor = None
        self.device = "cuda" if self._is_cuda_available() else "cpu"
        self._load_clip_model()
    
    def _is_cuda_available(self) -> bool:
        """Check if CUDA is available"""
        try:
            import torch
            return torch.cuda.is_available()
        except ImportError:
            return False
    
    def _load_clip_model(self):
        """Load CLIP model and processor"""
        try:
            from transformers import CLIPProcessor, CLIPModel
            
            print(f"🖼️  Loading CLIP model on {self.device}...")
            model_name = "openai/clip-vit-base-patch32"
            
            self.clip_model = CLIPModel.from_pretrained(model_name)
            self.clip_processor = CLIPProcessor.from_pretrained(model_name)
            
            if self.device == "cuda":
                self.clip_model = self.clip_model.to(self.device)
            
            print(f"✅ CLIP model loaded successfully on {self.device}")
            
        except ImportError as e:
            print(f"⚠️  CLIP not available: {e}")
            print("💡 Install with: pip install transformers torch torchvision")
            self.clip_model = None
            self.clip_processor = None
        except Exception as e:
            print(f"❌ Error loading CLIP model: {e}")
            self.clip_model = None
            self.clip_processor = None
    
    def generate_image_embedding(self, image_path: str) -> Union[List[float], None]:
        """Generate CLIP embedding for an image"""
        if not self.clip_model or not self.clip_processor:
            print("❌ CLIP model not available")
            return None
        
        try:
            # Load and preprocess image (handle SVG conversion)
            image = self._load_image_as_raster(image_path)
            if image is None:
                return None
            
            # Process image with CLIP
            inputs = self.clip_processor(images=image, return_tensors="pt")
            
            if self.device == "cuda":
                inputs = {k: v.to(self.device) for k, v in inputs.items()}
            
            # Generate embedding
            import torch
            with torch.no_grad():
                image_features = self.clip_model.get_image_features(**inputs)
                embedding = image_features.cpu().numpy().flatten().tolist()
            
            return embedding
            
        except Exception as e:
            print(f"❌ Error generating image embedding: {e}")
            return None
    
    def _load_image_as_raster(self, image_path: str) -> Union[Image.Image, None]:
        """Load image as raster format, converting SVG if necessary"""
        try:
            file_ext = os.path.splitext(image_path)[1].lower()
            
            if file_ext == '.svg':
                # Convert SVG to raster image
                return self._convert_svg_to_raster(image_path)
            else:
                # Load regular raster image
                return Image.open(image_path).convert('RGB')
                
        except Exception as e:
            print(f"❌ Error loading image {image_path}: {e}")
            return None
    
    def _convert_svg_to_raster(self, svg_path: str, size: Tuple[int, int] = (1024, 1024)) -> Union[Image.Image, None]:
        """Convert SVG to raster image using cairosvg"""
        try:
            # Read SVG file
            with open(svg_path, 'rb') as svg_file:
                svg_data = svg_file.read()
            
            # Convert SVG to PNG bytes
            png_data = cairosvg.svg2png(
                bytestring=svg_data,
                output_width=size[0],
                output_height=size[1],
                dpi=96
            )
            
            # Convert PNG bytes to PIL Image
            image = Image.open(io.BytesIO(png_data)).convert('RGB')
            
            print(f"  - Converted SVG to raster: {size[0]}x{size[1]} pixels")
            return image
            
        except ImportError:
            print("⚠️  cairosvg not available. Install with: pip install cairosvg")
            return None
        except Exception as e:
            print(f"❌ Error converting SVG: {e}")
            return None
    
    def generate_text_embedding(self, text: str) -> Union[List[float], None]:
        """Generate CLIP text embedding"""
        if not self.clip_model or not self.clip_processor:
            print("❌ CLIP model not available")
            return None
        
        try:
            # Process text with CLIP
            inputs = self.clip_processor(text=text, return_tensors="pt", padding=True)
            
            if self.device == "cuda":
                inputs = {k: v.to(self.device) for k, v in inputs.items()}
            
            # Generate embedding
            import torch
            with torch.no_grad():
                text_features = self.clip_model.get_text_features(**inputs)
                embedding = text_features.cpu().numpy().flatten().tolist()
            
            return embedding
            
        except Exception as e:
            print(f"❌ Error generating text embedding: {e}")
            return None
    
    def get_image_description(self, image_path: str) -> str:
        """Generate a comprehensive description of the image"""
        if not self.clip_model or not self.clip_processor:
            return f"Image file: {os.path.basename(image_path)}"
        
        try:
            file_ext = os.path.splitext(image_path)[1].lower()
            file_size = os.path.getsize(image_path)
            
            # Handle SVG files specially
            if file_ext == '.svg':
                description = f"Image: {os.path.basename(image_path)}"
                description += f"\nFormat: SVG (Vector Graphics)"
                description += f"\nFile size: {self._format_file_size(file_size)}"
                description += f"\nNote: Converted to raster for CLIP processing"
                
                # Try to get SVG dimensions if possible
                try:
                    import xml.etree.ElementTree as ET
                    tree = ET.parse(image_path)
                    root = tree.getroot()
                    
                    # Get SVG dimensions from viewBox or width/height attributes
                    viewbox = root.get('viewBox')
                    width = root.get('width')
                    height = root.get('height')
                    
                    if viewbox:
                        parts = viewbox.split()
                        if len(parts) >= 4:
                            description += f"\nOriginal dimensions: {parts[2]}x{parts[3]} units"
                    elif width and height:
                        description += f"\nOriginal dimensions: {width}x{height} units"
                        
                except Exception:
                    description += "\nOriginal dimensions: Unknown"
                
                return description
            
            # Handle raster images
            else:
                # Load image
                image = Image.open(image_path).convert('RGB')
                
                # Get image dimensions and basic info
                width, height = image.size
                
                # Try to extract some basic features
                description = f"Image: {os.path.basename(image_path)}"
                description += f"\nDimensions: {width}x{height} pixels"
                description += f"\nFile size: {self._format_file_size(file_size)}"
                description += f"\nFormat: {image.format}"
                
                # Add color information
                if image.mode == 'RGB':
                    # Sample some pixels for color analysis
                    pixels = list(image.getdata())
                    sample_size = min(1000, len(pixels))
                    sample_pixels = pixels[::len(pixels)//sample_size][:sample_size]
                    
                    # Calculate average RGB values
                    avg_r = sum(p[0] for p in sample_pixels) / len(sample_pixels)
                    avg_g = sum(p[1] for p in sample_pixels) / len(sample_pixels)
                    avg_b = sum(p[2] for p in sample_pixels) / len(sample_pixels)
                    
                    description += f"\nAverage color: RGB({int(avg_r)}, {int(avg_g)}, {int(avg_b)})"
                    
                    # Add dominant colors analysis
                    dominant_colors = self._analyze_dominant_colors(image)
                    if dominant_colors:
                        description += f"\nDominant colors: {', '.join(dominant_colors)}"
                
                return description
            
        except Exception as e:
            return f"Image file: {os.path.basename(image_path)} (Error: {e})"
    
    def _analyze_dominant_colors(self, image: Image.Image, num_colors=5) -> List[str]:
        """Analyze dominant colors in an image"""
        try:
            # Resize image for faster processing
            small_image = image.resize((150, 150))
            pixels = list(small_image.getdata())
            
            # Group similar colors
            color_groups = {}
            for pixel in pixels:
                # Round RGB values to group similar colors
                rounded = (pixel[0]//25*25, pixel[1]//25*25, pixel[2]//25*25)
                if rounded in color_groups:
                    color_groups[rounded] += 1
                else:
                    color_groups[rounded] = 1
            
            # Sort by frequency and get top colors
            sorted_colors = sorted(color_groups.items(), key=lambda x: x[1], reverse=True)
            dominant_colors = []
            
            for (r, g, b), count in sorted_colors[:num_colors]:
                # Convert to color names
                color_name = self._rgb_to_color_name(r, g, b)
                dominant_colors.append(color_name)
            
            return dominant_colors
            
        except Exception:
            return []
    
    def _rgb_to_color_name(self, r: int, g: int, b: int) -> str:
        """Convert RGB values to color names"""
        # Basic color mapping
        if r > 200 and g > 200 and b > 200:
            return "White"
        elif r < 50 and g < 50 and b < 50:
            return "Black"
        elif r > 200 and g < 100 and b < 100:
            return "Red"
        elif r < 100 and g > 200 and b < 100:
            return "Green"
        elif r < 100 and g < 100 and b > 200:
            return "Blue"
        elif r > 200 and g > 200 and b < 100:
            return "Yellow"
        elif r > 200 and g < 100 and b > 200:
            return "Magenta"
        elif r < 100 and g > 200 and b > 200:
            return "Cyan"
        elif r > 150 and g > 100 and b < 100:
            return "Orange"
        elif r > 100 and g > 150 and b < 100:
            return "Lime"
        elif r < 100 and g > 150 and b > 150:
            return "Teal"
        elif r > 150 and g < 100 and b > 150:
            return "Purple"
        else:
            return f"RGB({r},{g},{b})"
    
    def process_image_directory(self, directory_path: str) -> List[dict]:
        """Process all images in a directory and return analysis results"""
        results = []
        
        if not os.path.exists(directory_path):
            print(f"❌ Directory '{directory_path}' does not exist!")
            return results
        
        # Get all image files
        image_extensions = ('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.webp', '.svg')
        image_files = [f for f in os.listdir(directory_path) 
                      if f.lower().endswith(image_extensions)]
        
        if not image_files:
            print(f"❌ No image files found in '{directory_path}'")
            return results
        
        print(f"🖼️  Processing {len(image_files)} images in '{directory_path}'...")
        
        for i, filename in enumerate(image_files, 1):
            file_path = os.path.join(directory_path, filename)
            print(f"  [{i}/{len(image_files)}] Processing: {filename}")
            
            try:
                # Get image description
                description = self.get_image_description(file_path)
                
                # Generate CLIP embedding
                embedding = self.generate_image_embedding(file_path)
                
                result = {
                    'filename': filename,
                    'path': file_path,
                    'description': description,
                    'has_embedding': embedding is not None,
                    'embedding_dimensions': len(embedding) if embedding else 0
                }
                
                results.append(result)
                
                if embedding:
                    print(f"    ✅ CLIP embedding generated ({len(embedding)} dimensions)")
                else:
                    print(f"    ⚠️  No CLIP embedding generated")
                    
            except Exception as e:
                print(f"    ❌ Error processing {filename}: {e}")
                results.append({
                    'filename': filename,
                    'path': file_path,
                    'description': f"Error: {e}",
                    'has_embedding': False,
                    'embedding_dimensions': 0
                })
        
        print(f"🎉 Image directory processing completed!")
        print(f"📊 Summary: {len(results)} images processed")
        
        return results
    
    def _format_file_size(self, size_bytes: int) -> str:
        """Format file size in human-readable format"""
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size_bytes < 1024.0:
                return f"{size_bytes:.1f} {unit}"
            size_bytes /= 1024.0
        return f"{size_bytes:.1f} TB"

class FileLoader:
    @staticmethod
    def detect_encoding(file_path: str) -> str:
        """Detect file encoding"""
        with open(file_path, 'rb') as file:
            raw_data = file.read()
            result = chardet.detect(raw_data)
            return result['encoding']

    @staticmethod
    def load_excel(file_path: str) -> List[Document]:
        """Load Excel file and convert to documents"""
        try:
            # Read Excel file with all sheets
            excel_file = pd.ExcelFile(file_path)
            documents = []
            
            # Process each sheet
            for sheet_name in excel_file.sheet_names:
                # Read the sheet
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # Convert DataFrame to string representation
                excel_text = f"Sheet Name: {sheet_name}\n\n"
                excel_text += df.to_string(index=False)
                
                # Convert column names to string for metadata
                columns_str = ", ".join(df.columns.tolist())
                
                # Create a document with metadata
                doc = Document(
                    page_content=excel_text,
                    metadata={
                        "source": file_path,
                        "type": "excel",
                        "sheet_name": sheet_name,
                        "rows": str(len(df)),  # Convert to string
                        "columns": str(len(df.columns)),  # Convert to string
                        "columns_list": columns_str  # Store as comma-separated string
                    }
                )
                documents.append(doc)
                print(f"  - Loaded sheet: {sheet_name} with {len(df)} rows and {len(df.columns)} columns")
                print(f"    Columns: {columns_str}")
            
            return documents
        except Exception as e:
            print(f"Error loading Excel file {file_path}: {e}")
            return []

    @staticmethod
    def load_csv(file_path: str) -> List[Document]:
        """Load CSV file and convert to documents"""
        try:
            # Detect encoding
            encoding = FileLoader.detect_encoding(file_path)
            
            # Read CSV file
            df = pd.read_csv(file_path, encoding=encoding)
            
            # Convert DataFrame to string representation
            csv_text = df.to_string(index=False)
            
            # Create a document with metadata
            doc = Document(
                page_content=csv_text,
                metadata={
                    "source": file_path,
                    "type": "csv",
                    "rows": len(df),
                    "columns": len(df.columns),
                    "columns_list": list(df.columns)
                }
            )
            return [doc]
        except Exception as e:
            print(f"Error loading CSV file {file_path}: {e}")
            return []

    @staticmethod
    def load_word(file_path: str) -> List[Document]:
        """Load Word document and convert to documents"""
        try:
            # Read Word document
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Create a document with metadata
            return [Document(
                page_content=text,
                metadata={
                    "source": file_path,
                    "type": "word",
                    "paragraphs": len(doc.paragraphs)
                }
            )]
        except Exception as e:
            print(f"Error loading Word file {file_path}: {e}")
            return []

    @staticmethod
    def load_pdf(file_path: str) -> List[Document]:
        """Load PDF file and convert to documents"""
        try:
            # Read PDF file
            reader = PdfReader(file_path)
            
            # Extract text from all pages
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            
            # Create a document with metadata
            return [Document(
                page_content=text,
                metadata={
                    "source": file_path,
                    "type": "pdf",
                    "pages": len(reader.pages)
                }
            )]
        except Exception as e:
            print(f"Error loading PDF file {file_path}: {e}")
            return []
    
    @staticmethod
    def load_image(file_path: str, image_loader: 'ImageLoader' = None) -> List[Document]:
        """Load image file and convert to documents with CLIP embeddings"""
        try:
            # Create image loader if not provided
            if image_loader is None:
                image_loader = ImageLoader()
            
            # Get image description
            image_description = image_loader.get_image_description(file_path)
            
            # Generate CLIP embedding if available
            clip_embedding = None
            if image_loader.clip_model:
                clip_embedding = image_loader.generate_image_embedding(file_path)
                if clip_embedding:
                    print(f"  - Generated CLIP embedding: {len(clip_embedding)} dimensions")
            
            # Create metadata
            metadata = {
                "source": file_path,
                "type": "image",
                "has_clip_embedding": clip_embedding is not None
            }
            
            # Add CLIP embedding to metadata if available
            if clip_embedding:
                metadata["clip_embedding"] = clip_embedding
                metadata["embedding_dimensions"] = len(clip_embedding)
            
            # Create document
            return [Document(
                page_content=image_description,
                metadata=metadata
            )]
            
        except Exception as e:
            print(f"Error loading image file {file_path}: {e}")
            return []

class VectorDBChat:
    def __init__(self, model_type: str = 'gemini', model_name: str = None, 
                 temperature: float = 0.7, max_tokens: int = None, 
                 enable_web_search: bool = True, chunk_size: int = 1000, 
                 chunk_overlap: int = 200, verbose_web_search: bool = False):
        
        # Initialize AI model factory
        self.model_factory = AIModelFactory(model_type, model_name, temperature, max_tokens)
        self.model_type = model_type
        self.model_name = self.model_factory.model_name
        self.temperature = temperature
        self.max_tokens = max_tokens
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.verbose_web_search = verbose_web_search
        
        # Initialize the chat model
        self.chat = self.model_factory.create_chat_model()

        # Initialize embeddings
        self.embeddings = self.model_factory.create_embedding_model()

        # Initialize vector store
        self.vector_store = None
        self.initialize_vector_store()
        
        # Initialize web search (optional)
        self.web_search = WebSearch() if enable_web_search else None
        
        # Initialize voice manager
        try:
            self.voice_manager = VoiceManager(enable_voice=True)
        except ImportError as e:
            print(f"⚠️  Voice features disabled (missing packages): {e}")
            print("💡 Install with: pip install SpeechRecognition pyttsx3 pyaudio")
            self.voice_manager = None
        except Exception as e:
            print(f"⚠️  Voice features disabled: {e}")
            self.voice_manager = None

    def initialize_vector_store(self):
        """Initialize or load the vector store"""
        try:
            # Create new vector store with persistence
            self.vector_store = Chroma(
                persist_directory="./chroma_data",
                embedding_function=self.embeddings
            )
            print("Initialized vector store")
        except Exception as e:
            print(f"Error initializing vector store: {e}")
            raise
    
    def clear_vector_store(self):
        """Clear the entire vector store and recreate it"""
        try:
            print("🗑️  Clearing vector store...")
            
            # Delete the chroma_data directory
            import shutil
            if os.path.exists("./chroma_data"):
                shutil.rmtree("./chroma_data")
                print("✅ Deleted existing chroma_data directory")
            
            # Reinitialize the vector store
            self.initialize_vector_store()
            print("✅ Vector store cleared and reinitialized successfully")
            
        except Exception as e:
            print(f"❌ Error clearing vector store: {e}")
            raise
    
    def get_vector_store_info(self):
        """Get information about the current vector store"""
        try:
            if self.vector_store is None:
                print("❌ Vector store not initialized")
                return
            
            # Get collection info
            collection = self.vector_store._collection
            if collection is None:
                print("ℹ️  Vector store is empty (no collections)")
                return
            
            count = collection.count()
            print(f"📊 Vector store contains {count} documents")
            
            # Check if chroma_data directory exists
            if os.path.exists("./chroma_data"):
                dir_size = self._get_directory_size("./chroma_data")
                print(f"💾 Database size: {dir_size}")
            
        except Exception as e:
            print(f"❌ Error getting vector store info: {e}")
    
    def _get_directory_size(self, path):
        """Calculate directory size in human-readable format"""
        total_size = 0
        try:
            for dirpath, dirnames, filenames in os.walk(path):
                for filename in filenames:
                    filepath = os.path.join(dirpath, filename)
                    if os.path.exists(filepath):
                        total_size += os.path.getsize(filepath)
            
            # Convert to human-readable format
            for unit in ['B', 'KB', 'MB', 'GB']:
                if total_size < 1024.0:
                    return f"{total_size:.1f} {unit}"
                total_size /= 1024.0
            return f"{total_size:.1f} TB"
        except:
            return "Unknown"

    def add_documents(self, directory_path="./documents"):
        """Add documents to the vector store"""
        try:
            documents = []
            
            print("\nScanning for files in:", directory_path)
            
            # Check if directory exists and has files
            if not os.path.exists(directory_path):
                print(f"❌ Directory '{directory_path}' does not exist!")
                return
            
            files = os.listdir(directory_path)
            if not files:
                print(f"❌ No files found in '{directory_path}' directory!")
                print("💡 Please add some documents (Excel, CSV, Word, PDF, or text files) to the 'documents' folder.")
                return
            
            # Load text files
            print("\nLoading text files...")
            text_loader = DirectoryLoader(directory_path, glob="**/*.txt")
            text_docs = text_loader.load()
            documents.extend(text_docs)
            print(f"Found {len(text_docs)} text files")
            
            # Initialize image loader for CLIP embeddings
            try:
                image_loader = ImageLoader()
                print("  - CLIP image processing enabled")
            except Exception as e:
                print(f"  - CLIP image processing disabled: {e}")
                image_loader = None
            
            # Process other file types
            print("\nLoading other file types...")
            for filename in files:
                file_path = os.path.join(directory_path, filename)
                print(f"Processing: {filename}")
                
                if filename.endswith(('.xlsx', '.xls')):
                    docs = FileLoader.load_excel(file_path)
                    documents.extend(docs)
                    print(f"  - Loaded Excel file: {filename}")
                elif filename.endswith('.csv'):
                    docs = FileLoader.load_csv(file_path)
                    documents.extend(docs)
                    print(f"  - Loaded CSV file: {filename}")
                elif filename.endswith(('.doc', '.docx')):
                    docs = FileLoader.load_word(file_path)
                    documents.extend(docs)
                    print(f"  - Loaded Word file: {filename}")
                elif filename.endswith('.pdf'):
                    docs = FileLoader.load_pdf(file_path)
                    documents.extend(docs)
                    print(f"  - Loaded PDF file: {filename}")
                elif filename.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.webp', '.svg')):
                    if image_loader:
                        try:
                            docs = FileLoader.load_image(file_path, image_loader)
                            documents.extend(docs)
                            print(f"  - Loaded image file: {filename}")
                        except Exception as e:
                            print(f"  - Failed to load image {filename}: {e}")
                    else:
                        print(f"  - Skipped image {filename} (CLIP not available)")
                elif filename.endswith('.txt'):
                    # Skip .txt files as they're already loaded by DirectoryLoader
                    continue
                else:
                    print(f"  - Skipped unsupported file type: {filename}")

            if not documents:
                print("\n⚠️  No supported documents found to process!")
                print("💡 Supported formats: Excel (.xlsx, .xls), CSV (.csv), Word (.doc, .docx), PDF (.pdf), Text (.txt), Images (.png, .jpg, .jpeg, .gif, .bmp, .tiff, .webp, .svg)")
                print("💡 You can still chat with the AI, but it won't have document context.")
                # Don't return here - continue to chat interface
            else:
                print("\nSplitting documents into chunks...")
                # Split documents into chunks
                text_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=self.chunk_size,
                    chunk_overlap=self.chunk_overlap
                )
                splits = text_splitter.split_documents(documents)
                print(f"Created {len(splits)} chunks from {len(documents)} documents")

                # Add to vector store
                print("\nAdding chunks to vector store...")
                self.vector_store.add_documents(splits)
                print(f"✅ Successfully added {len(splits)} chunks to vector store")
                
                # Print summary of processed files
                print("\n📋 Processed files summary:")
                for doc in documents:
                    print(f"- {doc.metadata.get('source', 'Unknown')} ({doc.metadata.get('type', 'text')})")
                    if doc.metadata.get('type') == 'excel':
                        print(f"  Sheet: {doc.metadata.get('sheet_name')}")
                        print(f"  Columns: {doc.metadata.get('columns_list')}")



        except Exception as e:
            print(f"❌ Error adding documents: {e}")
            import traceback
            print(traceback.format_exc())

    def chat_with_context(self):
        """Chat with AI using vector store context and web search when needed"""
        colored_print(f"Welcome to the AI Chat with {self.model_factory.model_config['name']}!", Fore.CYAN, Style.BRIGHT)
        colored_print(f"Model: {self.model_name}", Fore.CYAN)
        colored_print(f"Temperature: {self.temperature}", Fore.CYAN)
        if self.max_tokens:
            colored_print(f"Max Tokens: {self.max_tokens}", Fore.CYAN)
        colored_print("Type 'help' to see all available commands.", Fore.WHITE)
        colored_print("Type 'quit' or 'exit' to end the conversation.", Fore.WHITE)
        colored_print("Type 'clear' to clear chat history, or 'clear db' to manage database.", Fore.WHITE)
        colored_print("Type 'search image <description>' to find images.", Fore.WHITE)
        colored_print("Type 'image summary' to see image database stats.", Fore.WHITE)
        colored_print("Type 'history' to see conversation summary.", Fore.WHITE)
        colored_print("Type your message and press Enter to chat with the AI.", Fore.WHITE)
        colored_print("You can ask questions about the data in your documents (Excel, CSV, Word, PDF, text files, and images).", Fore.WHITE)
        if self.web_search:
            colored_print("If the answer isn't in your documents, I can search the web (with your permission).", Fore.WHITE)
        else:
            colored_print("Web search is disabled.", Fore.WHITE)
        colored_print("-" * 50, Fore.CYAN)

        # Initialize conversation history
        conversation_history = []
        
        # Add system message to establish context
        system_message = f"""You are an AI assistant with access to a knowledge base. 
        You can help answer questions about documents, data, and general topics.
        Model: {self.model_name}
        Temperature: {self.temperature}
        Web Search: {'Enabled' if self.web_search else 'Disabled'}
        
        Please provide helpful, accurate, and context-aware responses.
        If you're referencing previous parts of the conversation, acknowledge them appropriately."""
        
        conversation_history.append({"role": "system", "content": system_message})
        
        colored_print(f"\n💬 Chat session started! Type 'help' for available commands.", Fore.GREEN)
        colored_print(f"📝 Conversation history will be maintained throughout this session.", Fore.GREEN)

        while True:
            # Get user input
            user_input = input(f"\n{Fore.BLUE}You: {Style.RESET_ALL}").strip()
            
            # Check if user wants to quit or clear database
            if user_input.lower() in ['quit', 'exit']:
                colored_print("\nGoodbye! Thanks for chatting!", Fore.GREEN, Style.BRIGHT)
                break
            elif user_input.lower() in ['clear', 'clear chat']:
                colored_print("\n🗑️  Clearing chat history...", Fore.YELLOW)
                conversation_history = [conversation_history[0]]  # Keep system message
                colored_print("✅ Chat history cleared!", Fore.GREEN)
                continue
            elif user_input.lower() in ['clear db', 'clear database']:
                colored_print("\n🗑️  Database Management Options:", Fore.YELLOW)
                colored_print("1. Clear vector database and reload documents", Fore.WHITE)
                colored_print("2. Just view database info", Fore.WHITE)
                colored_print("3. Cancel", Fore.WHITE)
                
                clear_choice = input(f"{Fore.CYAN}Enter choice (1-3): {Style.RESET_ALL}").strip()
                
                if clear_choice == "1":
                    colored_print("\n⚠️  WARNING: This will delete ALL existing data!", Fore.RED)
                    confirm = input(f"{Fore.RED}Are you sure? Type 'YES' to confirm: {Style.RESET_ALL}").strip()
                    if confirm == "YES":
                        self.clear_vector_store()
                        colored_print("\n📚 Reloading documents...", Fore.CYAN)
                        self.add_documents()
                        colored_print("✅ Database cleared and reloaded successfully!", Fore.GREEN)
                    else:
                        colored_print("❌ Database clearing cancelled.", Fore.YELLOW)
                elif clear_choice == "2":
                    self.get_vector_store_info()
                elif clear_choice == "3":
                    colored_print("✅ Cancelled.", Fore.GREEN)
                else:
                    colored_print("❌ Invalid choice.", Fore.RED)
                
                continue
            elif user_input.lower() == 'help':
                colored_print("\n" + "="*50, Fore.CYAN)
                colored_print("📚 Available Commands:", Fore.YELLOW, Style.BRIGHT)
                colored_print("  help              - Show this help message", Fore.WHITE)
                colored_print("  clear             - Clear chat history", Fore.WHITE)
                colored_print("  clear db          - Manage database", Fore.WHITE)
                colored_print("  search image <desc> - Find images by description", Fore.WHITE)
                colored_print("  image summary     - Show image database stats", Fore.WHITE)
                colored_print("  process images <dir> - Process all images in directory", Fore.WHITE)
                colored_print("  history           - Show conversation summary", Fore.WHITE)
                colored_print("  voice mode        - Switch to voice interaction", Fore.WHITE)
                colored_print("  speak             - Text-to-speech mode", Fore.WHITE)
                colored_print("  quit/exit         - End conversation", Fore.WHITE)
                colored_print("\n🎤 Voice Commands:", Fore.YELLOW, Style.BRIGHT)
                colored_print("  Say 'hey ai' to wake up in voice mode", Fore.WHITE)
                colored_print("  Say 'exit voice' to return to text mode", Fore.WHITE)
                colored_print("="*50, Fore.CYAN)
                continue
            elif user_input.lower() == 'history':
                colored_print("\n" + "="*50, Fore.CYAN)
                colored_print("📝 Conversation Summary:", Fore.YELLOW, Style.BRIGHT)
                user_messages = [msg for msg in conversation_history if msg["role"] == "user"]
                ai_messages = [msg for msg in conversation_history if msg["role"] == "assistant"]
                colored_print(f"  💬 User messages: {len(user_messages)}", Fore.WHITE)
                colored_print(f"  🤖 AI responses: {len(ai_messages)}", Fore.WHITE)
                colored_print(f"  📊 Total exchanges: {len(user_messages)}", Fore.WHITE)
                if user_messages:
                    colored_print(f"  🎯 Last topic: {user_messages[-1]['content'][:100]}...", Fore.WHITE)
                colored_print("="*50, Fore.CYAN)
                continue
            elif user_input.lower().startswith('search image'):
                # Extract search query
                query = user_input.lower().replace('search image', '').strip()
                if query:
                    print(f"\n🔍 Searching for images: '{query}'")
                    self.search_images_by_description(query)
                else:
                    print("💡 Usage: search image <description>")
                    print("Example: search image cat")
                continue
            elif user_input.lower() in ['image summary', 'images summary', 'summary images']:
                print("\n" + "="*50)
                print(self.get_image_summary())
                print("="*50)
                continue
            elif user_input.lower().startswith('process images'):
                # Extract directory path
                parts = user_input.split(' ', 2)
                if len(parts) >= 3:
                    directory_path = parts[2].strip()
                    print(f"\n🖼️  Processing images in directory: {directory_path}")
                    
                    # Initialize image loader if not already done
                    if not hasattr(self, 'image_loader'):
                        try:
                            self.image_loader = ImageLoader()
                        except Exception as e:
                            print(f"❌ Image processing not available: {e}")
                            continue
                    
                    # Process the directory
                    results = self.image_loader.process_image_directory(directory_path)
                    
                    if results:
                        print(f"\n📊 Image Analysis Results:")
                        for result in results:
                            print(f"\n📁 {result['filename']}")
                            print(f"   Path: {result['path']}")
                            print(f"   CLIP: {'✅' if result['has_embedding'] else '❌'}")
                            if result['has_embedding']:
                                print(f"   Dimensions: {result['embedding_dimensions']}")
                            print(f"   Description: {result['description'][:200]}...")
                else:
                    print("💡 Usage: process images <directory_path>")
                    print("Example: process images ./my_images")
                continue
            elif user_input.lower() == 'voice mode':
                print("\n🎤 Switching to voice mode...")
                if hasattr(self, 'voice_manager'):
                    question = self.voice_manager.start_voice_mode()
                    if question:
                        # Process the voice question
                        print(f"\n🎯 Processing voice question: {question}")
                        # Add to conversation history and process
                        conversation_history.append({"role": "user", "content": question})
                        # Continue with normal processing...
                else:
                    print("❌ Voice features not available")
                continue
            elif user_input.lower() == 'speak':
                print("\n🎤 Text-to-speech mode activated!")
                print("💡 Type your message and I'll speak it out loud.")
                print("💡 Type 'exit speak' to return to normal mode.")
                
                while True:
                    speak_input = input("\n🎤 Speak: ").strip()
                    if speak_input.lower() == 'exit speak':
                        print("✅ Returning to normal mode")
                        break
                    
                    if hasattr(self, 'voice_manager'):
                        self.voice_manager.speak(speak_input)
                    else:
                        print("❌ Voice features not available")
                continue

            try:
                # Add user message to conversation history
                conversation_history.append({"role": "user", "content": user_input})
                
                # Search for relevant documents
                docs = self.vector_store.similarity_search(user_input, k=3)
                context = "\n".join([doc.page_content for doc in docs])

                # Create prompt with context and conversation history
                recent_history = conversation_history[-6:]  # Last 6 messages for context
                history_context = "\n".join([f"{msg['role'].title()}: {msg['content']}" for msg in recent_history[:-1]])
                
                prompt = f"""Context information from documents:
                ---------------------
                {context}
                ---------------------
                
                Recent conversation history:
                ---------------------
                {history_context}
                ---------------------
                
                Current user question: {user_input}
                
                Please provide a helpful response that:
                1. Answers the current question using available context
                2. References relevant parts of the conversation if applicable
                3. Maintains conversation flow and context
                4. If the question is about data in Excel, CSV, Word, or PDF files, provide specific insights
                5. If you cannot provide a complete answer, indicate this clearly
                
                Remember the conversation context and build upon previous exchanges appropriately."""

                # Create messages for AI model
                messages = []
                messages.append(HumanMessage(content=prompt))

                # Get AI response
                response = self.chat.invoke(messages)
                
                # Print AI response
                colored_print(f"\n{Fore.YELLOW}AI: {Fore.GREEN}{response.content}", Fore.GREEN)
                
                # Add AI response to conversation history
                conversation_history.append({"role": "assistant", "content": response.content})
                
                # Check if we should suggest web search (only if enabled)
                if self.web_search and self._should_suggest_web_search(response.content, user_input):
                    colored_print("\n" + "="*50, Fore.CYAN)
                    colored_print("🔍 Web search suggested! I found some information in your documents, but I might be able to provide a more complete answer.", Fore.WHITE)
                    colored_print("Type 'yes' to search the web, or press Enter to continue with current information.", Fore.WHITE)
                    
                    web_search_choice = input(f"{Fore.CYAN}Web search? (yes/Enter): {Style.RESET_ALL}").strip().lower()
                    
                    if web_search_choice == 'yes':
                        self._perform_web_search(user_input, conversation_history, messages)
                
                # Additional proactive web search for specific question types
                elif self.web_search and self._is_likely_web_search_question(user_input):
                    colored_print("\n" + "="*50, Fore.CYAN)
                    colored_print("🌐 This looks like a question that might benefit from web search!", Fore.WHITE)
                    colored_print("Type 'yes' to search the web, or press Enter to continue with current information.", Fore.WHITE)
                    
                    web_search_choice = input(f"{Fore.CYAN}Web search? (yes/Enter): {Style.RESET_ALL}").strip().lower()
                    
                    if web_search_choice == 'yes':
                        self._perform_web_search(user_input, conversation_history, messages)

            except Exception as e:
                colored_print(f"\nAn error occurred: {e}", Fore.RED)
                colored_print("Please try again.", Fore.WHITE)
    
    def _should_suggest_web_search(self, ai_response: str, user_question: str) -> bool:
        """Determine if web search should be suggested"""
        # Keywords that suggest incomplete information
        incomplete_indicators = [
            "i don't have enough information",
            "i cannot provide",
            "i don't have access to",
            "based on the available information",
            "limited information",
            "i'm not sure",
            "i don't know",
            "i'm sorry",
            "there is no mention",
            "cannot answer",
            "no information available",
            "not found in",
            "no mention of",
            "therefore, i cannot"
        ]
        
        response_lower = ai_response.lower()
        question_lower = user_question.lower()
        
        if self.verbose_web_search:
            colored_print(f"🔍 Checking web search triggers for: '{user_question}'", Fore.CYAN)
        
        # Check if response suggests incomplete information
        for indicator in incomplete_indicators:
            if indicator in response_lower:
                if self.verbose_web_search:
                    colored_print(f"  ✅ Triggered by response indicator: '{indicator}'", Fore.GREEN)
                return True
        
        # Check if question is about current events, recent information, or external topics
        external_topics = [
            "latest", "recent", "current", "news", "today", "yesterday",
            "weather", "stock", "price", "market", "covid", "pandemic"
        ]
        
        for topic in external_topics:
            if topic in question_lower:
                if self.verbose_web_search:
                    colored_print(f"  ✅ Triggered by external topic: '{topic}'", Fore.GREEN)
                return True
        
        # Check if question is asking about a person, entity, or specific thing
        # that might not be in local documents
        person_entity_indicators = [
            "who is", "what is", "tell me about", "who are", "what are",
            "who was", "what was", "who were", "what were"
        ]
        
        for indicator in person_entity_indicators:
            if indicator in question_lower:
                if self.verbose_web_search:
                    colored_print(f"  ✅ Triggered by person/entity indicator: '{indicator}'", Fore.GREEN)
                return True
        
        # Check if the response explicitly says it can't find information
        if any(phrase in response_lower for phrase in ["no mention", "cannot answer", "no information"]):
            if self.verbose_web_search:
                colored_print(f"  ✅ Triggered by explicit 'no information' response", Fore.GREEN)
            return True
        
        # Check if the question is asking about something specific that might be online
        # but not in local documents (like people, companies, current events, etc.)
        # Look for questions that ask "who is X" or "what is X" where X is likely a person/entity
        import re
        
        # Pattern to detect "who is X" or "what is X" questions
        who_what_pattern = r'\b(who|what)\s+is\s+([a-zA-Z0-9\s\-\.]+?)(?:\?|$|\.)'
        matches = re.findall(who_what_pattern, question_lower)
        
        if matches:
            if self.verbose_web_search:
                colored_print(f"  ✅ Triggered by who/what question pattern: {matches}", Fore.GREEN)
            # If it's asking about a specific person/entity, suggest web search
            return True
        
        if self.verbose_web_search:
            colored_print("  ❌ No web search triggers found", Fore.YELLOW)
        
        return False
    
    def _is_likely_web_search_question(self, question: str) -> bool:
        """Check if a question is likely to benefit from web search"""
        question_lower = question.lower()
        
        if self.verbose_web_search:
            colored_print(f"🌐 Checking proactive web search for: '{question}'", Fore.CYAN)
        
        # Questions about people, entities, or specific things
        person_entity_patterns = [
            r'\bwho\s+is\b',
            r'\bwhat\s+is\b', 
            r'\btell\s+me\s+about\b',
            r'\bwho\s+are\b',
            r'\bwhat\s+are\b'
        ]
        
        for pattern in person_entity_patterns:
            if re.search(pattern, question_lower):
                if self.verbose_web_search:
                    colored_print(f"  ✅ Triggered by person/entity pattern: '{pattern}'", Fore.GREEN)
                return True
        
        # Questions about current events, recent information
        current_event_patterns = [
            r'\blatest\b',
            r'\brecent\b', 
            r'\bcurrent\b',
            r'\bnews\b',
            r'\btoday\b',
            r'\byesterday\b',
            r'\bthis\s+week\b',
            r'\bthis\s+month\b'
        ]
        
        for pattern in current_event_patterns:
            if re.search(pattern, question_lower):
                if self.verbose_web_search:
                    colored_print(f"  ✅ Triggered by current event pattern: '{pattern}'", Fore.GREEN)
                return True
        
        # Questions about specific companies, products, technologies
        tech_company_patterns = [
            r'\bchatgpt\b',
            r'\bopenai\b',
            r'\banthropic\b',
            r'\bclaude\b',
            r'\bgemini\b',
            r'\bllama\b',
            r'\bgoogle\b',
            r'\bapple\b',
            r'\bmicrosoft\b'
        ]
        
        for pattern in tech_company_patterns:
            if re.search(pattern, question_lower):
                if self.verbose_web_search:
                    colored_print(f"  ✅ Triggered by tech/company pattern: '{pattern}'", Fore.GREEN)
                return True
        
        if self.verbose_web_search:
            colored_print("  ❌ No proactive web search triggers found", Fore.YELLOW)
        
        return False
    
    def _perform_web_search(self, query: str, conversation_history: List, messages: List):
        """Perform web search and provide enhanced response"""
        try:
            print("\n🔍 Searching the web...")
            
            # Perform web search
            search_results = self.web_search.search_web(query, max_results=3)
            
            if not search_results:
                print("❌ No web search results found.")
                return
            
            print(f"✅ Found {len(search_results)} web results:")
            for i, result in enumerate(search_results, 1):
                print(f"  {i}. {result['title']}")
                print(f"     {result['snippet'][:100]}...")
                print(f"     URL: {result['url']}")
                print()
            
            # Get additional content from the first result
            print("📖 Fetching detailed content from the first result...")
            detailed_content = self.web_search.get_web_content(search_results[0]['url'])
            
            # Create enhanced prompt with web content
            web_context = f"""
            Web Search Results for: {query}
            
            Top Result:
            Title: {search_results[0]['title']}
            URL: {result['url']}
            Content: {detailed_content}
            
            Additional Results:
            {chr(10).join([f"- {r['title']}: {r['snippet']}" for r in search_results[1:]])}
            """
            
            enhanced_prompt = f"""Based on the following web search results, please provide a comprehensive answer to: {query}

            Web Search Results:
            {web_context}
            
            Please synthesize the information from the web search with any previous context to provide a complete answer.
            Include relevant URLs and sources when appropriate.
            """
            
            # Get enhanced AI response
            enhanced_messages = [HumanMessage(content=enhanced_prompt)]
            enhanced_response = self.chat.invoke(enhanced_messages)
            
            colored_print("\n" + "="*50, Fore.CYAN)
            colored_print("🌐 Enhanced Answer (with web search):", Fore.MAGENTA, Style.BRIGHT)
            colored_print("="*50, Fore.CYAN)
            colored_print(enhanced_response.content, Fore.GREEN)
            colored_print("="*50, Fore.CYAN)
            
            # Add enhanced response to conversation history
            conversation_history.append({"role": "assistant", "content": enhanced_response.content})
            
            # Add enhanced response to message history
            messages.append(enhanced_response)
            
        except Exception as e:
            print(f"❌ Error during web search: {e}")
            print("Continuing with available information...")
    
    def search_images_by_description(self, query: str, k: int = 3) -> List[Document]:
        """Search for images by text description using CLIP"""
        try:
            # First try to find images in the vector store
            docs = self.vector_store.similarity_search(query, k=k)
            
            # Filter for image documents
            image_docs = [doc for doc in docs if doc.metadata.get('type') == 'image']
            
            if image_docs:
                print(f"🖼️  Found {len(image_docs)} relevant images:")
                for i, doc in enumerate(image_docs, 1):
                    print(f"  {i}. {doc.metadata.get('source', 'Unknown')}")
                    print(f"     {doc.page_content[:100]}...")
                return image_docs
            else:
                print("🖼️  No relevant images found in your database.")
                return []
                
        except Exception as e:
            print(f"❌ Error searching images: {e}")
            return []
    
    def get_image_summary(self) -> str:
        """Get a summary of all images in the database"""
        try:
            # Search for all image documents
            docs = self.vector_store.similarity_search("image", k=100)
            image_docs = [doc for doc in docs if doc.metadata.get('type') == 'image']
            
            if not image_docs:
                return "No images found in the database."
            
            summary = f"📊 Image Database Summary:\n"
            summary += f"Total images: {len(image_docs)}\n\n"
            
            # Group by format
            formats = {}
            for doc in image_docs:
                source = doc.metadata.get('source', 'Unknown')
                ext = os.path.splitext(source)[1].lower()
                formats[ext] = formats.get(ext, 0) + 1
            
            summary += "Formats:\n"
            for fmt, count in formats.items():
                summary += f"  {fmt}: {count} images\n"
            
            # Count CLIP embeddings
            clip_count = sum(1 for doc in image_docs if doc.metadata.get('has_clip_embedding', False))
            summary += f"\nCLIP embeddings: {clip_count}/{len(image_docs)} images"
            
            return summary
            
        except Exception as e:
            return f"Error getting image summary: {e}"

def main():
    # Parse command line arguments
    args = parse_arguments()
    
    # Create documents directory if it doesn't exist
    os.makedirs("./documents", exist_ok=True)
    
    # Show color status
    if COLORS_AVAILABLE:
        colored_print("🎨 Colored output enabled", Fore.GREEN)
    else:
        colored_print("⚠️  Colored output disabled - install colorama for better experience", Fore.YELLOW)
        colored_print("💡 Install with: pip install colorama", Fore.WHITE)
    
    colored_print("🤖 AI Chat with Vector DB and Web Search - Multi-Model Support", Fore.CYAN, Style.BRIGHT)
    colored_print("=" * 60, Fore.CYAN)
    colored_print(f"🎯 Selected Model: {AI_MODELS[args.model]['name']}", Fore.YELLOW)
    colored_print(f"🔧 Model Type: {args.model}", Fore.YELLOW)
    if args.model_name:
        colored_print(f"📝 Model Name: {args.model_name}", Fore.YELLOW)
    colored_print(f"🌡️  Temperature: {args.temperature}", Fore.YELLOW)
    if args.max_tokens:
        colored_print(f"🔢 Max Tokens: {args.max_tokens}", Fore.YELLOW)
    colored_print(f"🔍 Web Search: {'Enabled' if not args.no_web_search else 'Disabled'}", Fore.YELLOW)
    colored_print(f"📄 Chunk Size: {args.chunk_size}", Fore.YELLOW)
    colored_print(f"🔄 Chunk Overlap: {args.chunk_overlap}", Fore.YELLOW)
    colored_print("=" * 60, Fore.CYAN)
    
    try:
        # Initialize chat with selected model
        chat = VectorDBChat(
            model_type=args.model,
            model_name=args.model_name,
            temperature=args.temperature,
            max_tokens=args.max_tokens,
            enable_web_search=not args.no_web_search,
            chunk_size=args.chunk_size,
            chunk_overlap=args.chunk_overlap,
            verbose_web_search=args.verbose_web_search
        )
        
        # Check if database exists
        if os.path.exists("./chroma_data"):
            colored_print("\n📊 Database Status:", Fore.CYAN)
            chat.get_vector_store_info()
            colored_print("✅ Database found - ready for chat!", Fore.GREEN)
        else:
            colored_print("\n⚠️  No database found!", Fore.YELLOW)
            colored_print("💡 To build a database, run: python build_db.py --rebuild", Fore.WHITE)
            colored_print("💡 Or add documents to the 'documents' folder and run: python build_db.py --add-docs", Fore.WHITE)
            colored_print("💡 You can still chat, but without document context.", Fore.WHITE)
        
        # Start chat immediately
        colored_print("\n💬 Starting chat session...", Fore.GREEN)
        chat.chat_with_context()
        
    except ValueError as e:
        colored_print(f"❌ Configuration Error: {e}", Fore.RED)
        colored_print("💡 Please check your environment variables and model configuration.", Fore.WHITE)
    except ImportError as e:
        colored_print(f"❌ Import Error: {e}", Fore.RED)
        colored_print("💡 Please install the required packages for your selected model.", Fore.WHITE)
    except Exception as e:
        colored_print(f"❌ Unexpected Error: {e}", Fore.RED)
        colored_print("💡 Please check your configuration and try again.", Fore.WHITE)

if __name__ == "__main__":
    main() 