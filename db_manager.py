#!/usr/bin/env python3
"""
Database Manager for AI Chat with Vector Database
Handles database operations: create, clear, update, backup, restore
"""

import os
import shutil
import argparse
import time
from pathlib import Path
import json
from dotenv import load_dotenv

# LangChain imports
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import DirectoryLoader
from langchain_chroma import Chroma
from langchain.schema import Document

# AI Model imports
from langchain_google_genai import GoogleGenerativeAIEmbeddings

# File processing imports
import pandas as pd
import docx
import chardet
from PyPDF2 import PdfReader

class DatabaseManager:
    """Manages ChromaDB vector database operations"""
    
    def __init__(self, db_path="./chroma_data", documents_dir="./documents"):
        self.db_path = Path(db_path)
        self.documents_dir = Path(documents_dir)
        self.vector_store = None
        self.embeddings = None
        self._setup_default_embeddings()
    
    def _setup_default_embeddings(self):
        """Setup default embeddings if none provided"""
        if not self.embeddings:
            try:
                # Try to use Gemini embeddings as default
                load_dotenv()
                
                api_key = os.getenv('GEMINI_API_KEY')
                if api_key:
                    self.embeddings = GoogleGenerativeAIEmbeddings(
                        model="models/embedding-001",
                        google_api_key=api_key
                    )
                    print("✅ Using Gemini embeddings")
                    return
                
                # Try OpenAI as fallback
                api_key = os.getenv('OPENAI_API_KEY')
                if api_key:
                    from langchain_openai import OpenAIEmbeddings
                    self.embeddings = OpenAIEmbeddings(
                        model="text-embedding-3-small",
                        openai_api_key=api_key
                    )
                    print("✅ Using OpenAI embeddings")
                    return
                
                # Try Ollama as fallback
                try:
                    from langchain_community.embeddings import OllamaEmbeddings
                    self.embeddings = OllamaEmbeddings(model="llama3.2:3b")
                    print("✅ Using Ollama embeddings")
                    return
                except ImportError:
                    pass
                
                print("⚠️  No API keys found. Please set GEMINI_API_KEY or OPENAI_API_KEY in your .env file")
                
            except ImportError as e:
                print(f"⚠️  Required packages not installed: {e}")
                print("💡 Install with: pip install langchain-google-genai langchain-openai")
            except Exception as e:
                print(f"⚠️  Error setting up embeddings: {e}")
    
    def set_embeddings(self, embeddings):
        """Set embedding model for database operations"""
        self.embeddings = embeddings
    
    def create_database(self, force=False):
        """Create new database from documents"""
        if not self.embeddings:
            print("❌ Embeddings not set. Call set_embeddings() first.")
            print("💡 Make sure you have API keys in your .env file:")
            print("   - GEMINI_API_KEY for Google Gemini")
            print("   - OPENAI_API_KEY for OpenAI")
            print("   - Or have Ollama running locally")
            return False
        
        try:
            print("📚 Creating new database...")
            
            # Clear existing database if force=True
            if force and self.db_path.exists():
                self.clear_database(confirm=True)
            
            # Initialize vector store
            self.vector_store = Chroma(
                persist_directory=str(self.db_path),
                embedding_function=self.embeddings
            )
            
            # Load and process documents
            documents = self._load_documents()
            if not documents:
                print("⚠️  No documents found to process.")
                return False
            
            # Split into chunks
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000, chunk_overlap=200
            )
            chunks = text_splitter.split_documents(documents)
            
            # Add to database
            self.vector_store.add_documents(chunks)
            print(f"✅ Database created with {len(chunks)} chunks from {len(documents)} documents")
            return True
            
        except Exception as e:
            print(f"❌ Error creating database: {e}")
            return False
    
    def clear_database(self, confirm=False):
        """Clear entire database"""
        if not confirm:
            print("⚠️  WARNING: This will delete ALL database data!")
            response = input("Type 'YES' to confirm: ").strip()
            if response != "YES":
                print("❌ Operation cancelled.")
                return False
        
        try:
            print("🗑️  Clearing database...")
            
            # Close vector store
            if self.vector_store:
                try:
                    self.vector_store._client.reset()
                except:
                    pass
                self.vector_store = None
            
            # Delete database directory
            if self.db_path.exists():
                shutil.rmtree(self.db_path)
                print("✅ Database directory deleted")
            
            # Recreate empty directory
            self.db_path.mkdir(exist_ok=True)
            print("✅ Fresh database directory created")
            return True
            
        except Exception as e:
            print(f"❌ Error clearing database: {e}")
            return False
    
    def recreate_database(self, confirm=False):
        """Recreate database from scratch"""
        if not confirm:
            print("🔄 This will recreate the entire database!")
            response = input("Type 'YES' to confirm: ").strip()
            if response != "YES":
                print("❌ Operation cancelled.")
                return False
        
        try:
            print("🔄 Recreating database...")
            
            # Clear existing database
            if not self.clear_database(confirm=True):
                return False
            
            # Create new database
            if not self.create_database():
                return False
            
            print("✅ Database recreated successfully!")
            return True
            
        except Exception as e:
            print(f"❌ Error recreating database: {e}")
            return False
    
    def update_database(self):
        """Update database with new documents"""
        try:
            print("🔄 Updating database...")
            
            # Check for changes
            if not self._has_changes():
                print("✅ Database is up to date")
                return True
            
            # Recreate database to include all changes
            print("📚 Recreating database with all documents...")
            return self.recreate_database(confirm=True)
            
        except Exception as e:
            print(f"❌ Error updating database: {e}")
            return False
    
    def backup_database(self, backup_path=None):
        """Create database backup"""
        if not self.db_path.exists():
            print("❌ No database to backup!")
            return False
        
        try:
            if not backup_path:
                timestamp = time.strftime("%Y%m%d_%H%M%S")
                backup_path = f"./chroma_backup_{timestamp}"
            
            backup_path = Path(backup_path)
            print(f"💾 Creating backup at: {backup_path}")
            
            # Copy database
            if backup_path.exists():
                shutil.rmtree(backup_path)
            shutil.copytree(self.db_path, backup_path)
            
            # Create backup info
            info = {
                'backup_created': time.ctime(),
                'source_path': str(self.db_path),
                'backup_path': str(backup_path)
            }
            
            with open(backup_path / "backup_info.json", 'w') as f:
                json.dump(info, f, indent=2)
            
            print(f"✅ Backup created successfully!")
            return True
            
        except Exception as e:
            print(f"❌ Error creating backup: {e}")
            return False
    
    def restore_database(self, backup_path, confirm=False):
        """Restore database from backup"""
        backup_path = Path(backup_path)
        
        if not backup_path.exists():
            print(f"❌ Backup path '{backup_path}' does not exist!")
            return False
        
        if not confirm:
            print(f"⚠️  This will overwrite the current database!")
            response = input("Type 'YES' to confirm: ").strip()
            if response != "YES":
                print("❌ Operation cancelled.")
                return False
        
        try:
            print(f"🔄 Restoring from backup: {backup_path}")
            
            # Close current vector store
            if self.vector_store:
                try:
                    self.vector_store._client.reset()
                except:
                    pass
                self.vector_store = None
            
            # Remove current database
            if self.db_path.exists():
                shutil.rmtree(self.db_path)
            
            # Restore from backup
            shutil.copytree(backup_path, self.db_path)
            
            # Remove backup info from restored database
            backup_info = self.db_path / "backup_info.json"
            if backup_info.exists():
                backup_info.unlink()
            
            print("✅ Database restored successfully!")
            return True
            
        except Exception as e:
            print(f"❌ Error restoring database: {e}")
            return False
    
    def get_status(self):
        """Get database status information"""
        print("\n" + "="*50)
        print("🗄️  DATABASE STATUS")
        print("="*50)
        
        if not self.db_path.exists():
            print("❌ Database does not exist")
            return
        
        # Basic info
        print(f"📍 Path: {self.db_path}")
        print(f"💾 Size: {self._get_directory_size(self.db_path)}")
        
        # Check if database has data
        if (self.db_path / "chroma.sqlite3").exists():
            print("✅ Database files found")
            
            # Try to get document count
            try:
                if self.vector_store and self.vector_store._collection:
                    count = self.vector_store._collection.count()
                    print(f"📄 Documents: {count} chunks")
                else:
                    print("📄 Documents: Unknown (vector store not initialized)")
            except:
                print("📄 Documents: Unknown (error reading collection)")
        else:
            print("⚠️  Database files incomplete")
        
        # Show available backups
        backups = self._list_backups()
        if backups:
            print(f"\n💾 Available Backups: {len(backups)}")
            for i, backup in enumerate(backups[:3]):
                print(f"  {i+1}. {backup['created']}")
            if len(backups) > 3:
                print(f"  ... and {len(backups) - 3} more")
        
        print("="*50)
    
    def _load_documents(self):
        """Load all supported documents"""
        documents = []
        
        if not self.documents_dir.exists():
            print(f"❌ Documents directory '{self.documents_dir}' not found!")
            return documents
        
        files = list(self.documents_dir.iterdir())
        if not files:
            print(f"❌ No files found in '{self.documents_dir}'")
            return documents
        
        print(f"📁 Found {len(files)} files")
        
        # Load text files
        text_loader = DirectoryLoader(str(self.documents_dir), glob="**/*.txt")
        text_docs = text_loader.load()
        documents.extend(text_docs)
        print(f"📄 Loaded {len(text_docs)} text files")
        
        # Process other files
        for file_path in files:
            if file_path.is_file() and not file_path.suffix.lower() == '.txt':
                try:
                    docs = self._load_single_file(file_path)
                    if docs:
                        documents.extend(docs)
                        print(f"✅ Loaded: {file_path.name}")
                except Exception as e:
                    print(f"❌ Failed to load {file_path.name}: {e}")
        
        return documents
    
    def _load_single_file(self, file_path):
        """Load a single file based on type"""
        try:
            ext = file_path.suffix.lower()
            
            if ext in ['.xlsx', '.xls']:
                return self._load_excel(file_path)
            elif ext == '.csv':
                return self._load_csv(file_path)
            elif ext in ['.doc', '.docx']:
                return self._load_word(file_path)
            elif ext == '.pdf':
                return self._load_pdf(file_path)
            else:
                print(f"  ⚠️  Unsupported: {ext}")
                return []
                
        except Exception as e:
            print(f"  ❌ Error loading {file_path.name}: {e}")
            return []
    
    def _load_excel(self, file_path):
        """Load Excel file"""
        try:
            excel_file = pd.ExcelFile(file_path)
            docs = []
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                if not df.empty:
                    content = f"Sheet: {sheet_name}\n\n{df.to_string(index=False)}"
                    doc = Document(
                        page_content=content,
                        metadata={'source': str(file_path), 'type': 'excel', 'sheet': sheet_name}
                    )
                    docs.append(doc)
            
            return docs
        except Exception as e:
            print(f"    ❌ Excel error: {e}")
            return []
    
    def _load_csv(self, file_path):
        """Load CSV file"""
        try:
            # Detect encoding
            with open(file_path, 'rb') as f:
                result = chardet.detect(f.read())
                encoding = result['encoding']
            
            df = pd.read_csv(file_path, encoding=encoding)
            if not df.empty:
                content = df.to_string(index=False)
                doc = Document(
                    page_content=content,
                    metadata={'source': str(file_path), 'type': 'csv'}
                )
                return [doc]
        except Exception as e:
            print(f"    ❌ CSV error: {e}")
        return []
    
    def _load_word(self, file_path):
        """Load Word file"""
        try:
            doc = docx.Document(file_path)
            content = "\n".join([p.text for p in doc.paragraphs])
            if content.strip():
                doc_obj = Document(
                    page_content=content,
                    metadata={'source': str(file_path), 'type': 'word'}
                )
                return [doc_obj]
        except Exception as e:
            print(f"    ❌ Word error: {e}")
        return []
    
    def _load_pdf(self, file_path):
        """Load PDF file"""
        try:
            reader = PdfReader(file_path)
            content = ""
            for i, page in enumerate(reader.pages):
                content += f"Page {i+1}:\n{page.extract_text()}\n\n"
            
            if content.strip():
                doc = Document(
                    page_content=content,
                    metadata={'source': str(file_path), 'type': 'pdf', 'pages': len(reader.pages)}
                )
                return [doc]
        except Exception as e:
            print(f"    ❌ PDF error: {e}")
        return []
    
    def _has_changes(self):
        """Check if documents have changed since last database update"""
        # Simple implementation - always return True to trigger update
        # In a real implementation, you'd track file modification times
        return True
    
    def _list_backups(self):
        """List available backups"""
        backups = []
        current_dir = Path("./")
        
        try:
            for item in current_dir.iterdir():
                if item.is_dir() and item.name.startswith("chroma_backup_"):
                    info_file = item / "backup_info.json"
                    if info_file.exists():
                        try:
                            with open(info_file, 'r') as f:
                                info = json.load(f)
                            backups.append({
                                'path': str(item),
                                'created': info.get('backup_created', 'Unknown')
                            })
                        except:
                            backups.append({
                                'path': str(item),
                                'created': 'Unknown'
                            })
        except Exception as e:
            print(f"❌ Error listing backups: {e}")
        
        return sorted(backups, key=lambda x: x['created'], reverse=True)
    
    def _get_directory_size(self, path):
        """Get directory size in human-readable format"""
        total_size = 0
        try:
            for dirpath, dirnames, filenames in os.walk(path):
                for filename in filenames:
                    filepath = Path(dirpath) / filename
                    if filepath.exists():
                        total_size += filepath.stat().st_size
            
            # Convert to human-readable
            for unit in ['B', 'KB', 'MB', 'GB']:
                if total_size < 1024.0:
                    return f"{total_size:.1f} {unit}"
                total_size /= 1024.0
            return f"{total_size:.1f} TB"
        except:
            return "Unknown"


def main():
    """Command-line interface"""
    parser = argparse.ArgumentParser(
        description='Database Manager for AI Chat with Vector Database',
        epilog="""
Examples:
  python db_manager.py --status                    # Show status
  python db_manager.py --create                   # Create database
  python db_manager.py --recreate                 # Recreate database
  python db_manager.py --clear                    # Clear database
  python db_manager.py --update                   # Update database
  python db_manager.py --backup                   # Create backup
  python db_manager.py --restore ./backup_path    # Restore backup
        """
    )
    
    parser.add_argument('--status', action='store_true', help='Show database status')
    parser.add_argument('--create', action='store_true', help='Create new database')
    parser.add_argument('--recreate', action='store_true', help='Recreate database')
    parser.add_argument('--clear', action='store_true', help='Clear database')
    parser.add_argument('--update', action='store_true', help='Update database')
    parser.add_argument('--backup', action='store_true', help='Create backup')
    parser.add_argument('--restore', type=str, help='Restore from backup')
    parser.add_argument('--db-path', type=str, default='./chroma_data', help='Database path')
    parser.add_argument('--documents-dir', type=str, default='./documents', help='Documents directory')
    parser.add_argument('--force', action='store_true', help='Force operations without confirmation')
    
    args = parser.parse_args()
    
    # Create manager
    db_manager = DatabaseManager(
        db_path=args.db_path,
        documents_dir=args.documents_dir
    )
    
    # Handle commands
    if args.status:
        db_manager.get_status()
    
    elif args.create:
        if db_manager.create_database(force=args.force):
            print("✅ Database created successfully!")
            db_manager.get_status()
        else:
            print("❌ Failed to create database!")
    
    elif args.recreate:
        if db_manager.recreate_database(confirm=args.force):
            print("✅ Database recreated successfully!")
            db_manager.get_status()
        else:
            print("❌ Failed to recreate database!")
    
    elif args.clear:
        if db_manager.clear_database(confirm=args.force):
            print("✅ Database cleared successfully!")
            db_manager.get_status()
        else:
            print("❌ Failed to clear database!")
    
    elif args.update:
        if db_manager.update_database():
            print("✅ Database updated successfully!")
            db_manager.get_status()
        else:
            print("❌ Failed to update database!")
    
    elif args.backup:
        if db_manager.backup_database():
            print("✅ Backup created successfully!")
        else:
            print("❌ Failed to create backup!")
    
    elif args.restore:
        if db_manager.restore_database(args.restore, confirm=args.force):
            print("✅ Database restored successfully!")
            db_manager.get_status()
        else:
            print("❌ Failed to restore database!")
    
    else:
        # No arguments, show status
        db_manager.get_status()
        print("\n💡 Use --help to see all available commands")


if __name__ == "__main__":
    main()

